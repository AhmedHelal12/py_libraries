from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# doc = Document("academy_1.docx")
# paras = doc.paragraphs
# print(len(paras))

# for para in paras:
#     print(para.text)

# print(paras[0].text)
# print(paras[1].text)

# print(len(paras[2].runs))
# print(paras[2].runs[0].text)
# print(paras[2].runs[1].text)
# print(paras[2].runs[2].text)
# print(paras[2].runs[3].text)
# print(paras[2].runs[4].text)


doc = Document()


## add a paragraph and a subparagraph
p = doc.add_paragraph("This is the first paragraph.")
p.add_run("bold text ").bold = True
p.add_run("and ")
p.add_run("italic text ").italic = True

doc.add_paragraph("This is a paragraph with title attribute","Title")

doc.add_heading("This is a Title",0)
doc.add_heading("This is heading 1",1)
doc.add_heading("This is heading 2",2)
doc.add_heading("This is heading 3",3)
doc.add_heading("This is heading 4",4)
doc.add_heading("This is heading 5",5)
## ordered and unordered lists

doc.add_paragraph("This is first list bullet.",style="List Bullet")
doc.add_paragraph("This is second list bullet.",style="List Bullet")
doc.add_paragraph("This is first list number.",style="List Number")
doc.add_paragraph("This is second list number.",style="List Number")

## Alighnment
p = doc.add_paragraph("هذا مجرد اختبار بسيط ")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

### Image
doc.add_picture("ka.png",width=Inches(5),height=Inches(7))


### Style 

print(doc.paragraphs[0].style)
print(doc.paragraphs[2].style)

doc.paragraphs[2].style = doc.styles['Heading 4']
doc.paragraphs[2].style.delete()

## save the document
doc.save("new.docx")